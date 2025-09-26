# utils/file_utils.py

import os
import tempfile
import requests
import re
from typing import List, Tuple, Dict, Optional

from slack_sdk import WebClient

# For text extraction
from PyPDF2 import PdfReader
import docx
import openpyxl  # .xlsx
import xlrd      # .xls (ensure xlrd<2.0 for .xls support)
import pandas as pd
import difflib

# LangChain v0.2+
from langchain_core.documents import Document

# --- Column aliases + preferred order for "product profile" rendering ---
COL_ALIASES_PROFILE: Dict[str, str] = {
    "product name": "product_name",
    "product": "product_name",
    "name": "product_name",

    "other names this product is called (e.g. old names people still use)": "other_names",
    "aka": "other_names",
    "aliases": "other_names",

    "brief description of what product does for a client": "description",

    "product manager": "product_manager",
    "product mgmt vp": "product_mgmt_vp",
    "vp product mgmt": "product_mgmt_vp",
    "development vp": "development_vp",
    "support owner": "support_owner",
    "2nd line or first line owner": "second_line_owner",
    "support director": "support_director",
    "support planner": "support_planner",
    "pillar": "pillar",
    "mmt name (meeting where the cross functional team's meet and vote)": "mmt_name",
    "mmt name https://w3.ibm.com/w3publisher/software-support-community/new-page": "mmt_url",
}
PREFERRED_PROFILE_ORDER: List[str] = [
    "product_name",
    "other_names",
    "description",
    "product_manager",
    "product_mgmt_vp",
    "development_vp",
    "support_owner",
    "second_line_owner",
    "support_director",
    "support_planner",
    "pillar",
    "mmt_name",
    "mmt_url",
]

NOT_FOUND_MSG = "I couldn't find relevant information in the file."

# -----------------------------------------------------------------------------
# Path / download helpers
# -----------------------------------------------------------------------------

def sanitize_filename(fn: str) -> str:
    """
    Replace any character that is not alphanumeric, dot, hyphen, or underscore 
    with an underscore. This avoids slugify/Unicode issues.
    """
    return re.sub(r'[^A-Za-z0-9_.-]', '_', fn)

def _split_name_ext(name: str) -> Tuple[str, str]:
    base, ext = os.path.splitext(name)
    return sanitize_filename(base), ext  # ext includes leading dot

def download_slack_file(client: WebClient, file_info: dict, timeout: int = 25) -> str:
    """
    Given a Slack file_info dict (from the file_shared event),
    download the file to a temporary location and return local path.
    """
    url = file_info.get("url_private_download")
    if not url:
        raise RuntimeError("No url_private_download on file_info")

    headers = {"Authorization": f"Bearer {client.token}"}

    # Stream to avoid loading large files fully into memory
    with requests.get(url, headers=headers, stream=True, timeout=timeout) as r:
        if not r.ok:
            raise RuntimeError(f"Failed to download file: HTTP {r.status_code}")
        original_name = file_info.get("name") or "uploaded_file"
        safe_base, ext = _split_name_ext(original_name)
        tmp_path = os.path.join(tempfile.gettempdir(), safe_base + ext)

        with open(tmp_path, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)

    return tmp_path

def _is_not_found(s: str) -> bool:
    """
    Returns True if `s` is an explicit 'not found' style string from our handlers.
    Keep this in one place so we can reuse it everywhere.
    """
    if not s:
        return True
    t = s.strip().lower()
    return (
        t.startswith("i couldn't find") or
        t.startswith("i can’t find") or
        t.startswith("i can't find") or
        t.startswith("i couldnt find")
    )
# -----------------------------------------------------------------------------
# Extraction helpers
# -----------------------------------------------------------------------------

def extract_text_from_file(path: str) -> str:
    """
    Basic text extraction: PDF, DOCX, Excel (.xlsx/.xls), CSV/TSV, or plain text.
    Note: classic .doc is not supported by python-docx.
    """
    ext = path.lower().split(".")[-1]

    if ext == "pdf":
        text_parts: List[str] = []
        try:
            reader = PdfReader(path)
            for i, page in enumerate(reader.pages):
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    # per-page failure: continue
                    continue
        except Exception:
            return ""
        return "\n".join([t for t in text_parts if t])

    elif ext == "docx":
        try:
            d = docx.Document(path)
            return "\n".join(p.text for p in d.paragraphs)
        except Exception:
            return ""

    elif ext == "doc":
        # python-docx cannot parse .doc (binary). Return empty to avoid misleading output.
        return ""

    elif ext == "xlsx":
        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            text: List[str] = []
            for sheet in wb:
                for row in sheet.iter_rows():
                    row_text = [cell.value for cell in row if cell.value is not None]
                    if row_text:
                        text.append(" ".join(map(str, row_text)))
            return "\n".join(text)
        except Exception:
            return ""

    elif ext == "xls":
        try:
            wb = xlrd.open_workbook(path)
            text: List[str] = []
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    row_text = [str(cell.value) for cell in sheet.row(row_idx) if cell.value not in (None, "")]
                    if row_text:
                        text.append(" ".join(row_text))
            return "\n".join(text)
        except Exception:
            return ""

    elif ext in ("csv", "tsv"):
        sep = "\t" if ext == "tsv" else ","
        try:
            df = pd.read_csv(path, sep=sep, dtype=str, keep_default_na=False)
            return "\n".join([" ".join(map(str, row)) for row in df.fillna("").values.tolist()])
        except Exception:
            return ""

    else:
        # Try reading as plain text
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

def dataframe_to_documents(df: pd.DataFrame, file_name: str) -> List[Document]:
    docs: List[Document] = []
    cols = list(df.columns)
    for i, row in df.iterrows():
        content = "; ".join(f"{col}: {row[col]}" for col in cols)
        docs.append(Document(page_content=content, metadata={"row_index": int(i), "file_name": file_name}))
    return docs

def extract_excel_as_table(path: str) -> pd.DataFrame:
    ext = path.lower().split(".")[-1]
    if ext not in ("xlsx", "xls"):
        raise ValueError("Not an Excel file")

    engine = "openpyxl" if ext == "xlsx" else "xlrd"

    # Read raw to detect header row
    df_raw = pd.read_excel(path, header=None, engine=engine, dtype=str)
    header_row = 0
    for i, row in df_raw.iterrows():
        non_empty = sum(bool(str(cell).strip()) for cell in row)
        if non_empty >= 2:
            header_row = int(i)
            break

    df = pd.read_excel(path, header=header_row, engine=engine, dtype=str)
    # Clean up columns and rows
    df = df.loc[:, ~df.columns.astype(str).str.match(r'^Unnamed', na=False)]
    df = df.dropna(how='all').reset_index(drop=True)
    return df

# -----------------------------------------------------------------------------
# Normalization helpers
# -----------------------------------------------------------------------------

def _clean_text(s: str) -> str:
    import re as _re
    x = str(s).lower()
    x = x.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
    x = _re.sub(r"[-_/]+", " ", x)
    x = _re.sub(r"\s+", " ", x).strip()
    return x

def _clean_entity(entity: str) -> str:
    e = _clean_text(entity)
    e = re.sub(r"^(the\s+)?(product|person|employee|user)\s+", "", e)
    return e.strip(" '\"")

def _two_way_contains(series: pd.Series, needle: str) -> pd.Series:
    s = series.astype(str).str.lower().fillna("")
    needle = (needle or "").lower()
    return s.str.contains(re.escape(needle), regex=True, na=False) | s.apply(lambda x: bool(x) and x in needle)

def _role_tokens(text: str) -> set:
    import re as _re
    norm = _clean_text(text)
    toks = set(norm.split())
    if _re.search(r"\b(2nd|l2|level 2|second level)\b", norm):
        toks.update({"second", "line"})
    return toks

def resolve_role_column(columns: pd.Index | List[str], col_query: str, min_score: float = 60.0):
    q_tokens = _role_tokens(col_query)
    role_keywords = {"owner", "manager", "lead", "support", "vp", "director", "planner", "head"}

    best_col, best_score = None, -1.0
    for c in columns:
        c_norm = _clean_text(c)
        c_tokens = _role_tokens(c_norm)

        common = len(q_tokens & c_tokens)
        coverage = common / max(1, len(q_tokens))
        score = 70.0 * coverage

        if any(k in q_tokens and k in c_tokens for k in role_keywords):
            score += 25.0
        if c_norm == _clean_text(col_query):
            score += 50.0

        if ("product" in c_tokens or "name" in c_tokens) and not (("product" in q_tokens) or ("name" in q_tokens)):
            score -= 100.0
        if "planner" in c_tokens and "planner" not in q_tokens:
            score -= 20.0

        if score > best_score:
            best_score, best_col = score, c

    if best_score < min_score or best_col is None:
        return (None, best_score)

    best_tokens = _role_tokens(best_col)
    if not any(k in best_tokens for k in role_keywords):
        return (None, best_score)

    return (best_col, best_score)

def _normalize_dashes(s: str) -> str:
    return s.replace("–", "-").replace("—", "-")

# -----------------------------------------------------------------------------
# Column mapping / product helpers
# -----------------------------------------------------------------------------

def _map_columns_profile(df: pd.DataFrame) -> dict:
    mapping: Dict[str, str] = {}
    for col in df.columns:
        key = _clean_text(str(col).strip().lower())
        if key in COL_ALIASES_PROFILE:
            mapping[col] = COL_ALIASES_PROFILE[key]
            continue
        best = difflib.get_close_matches(key, COL_ALIASES_PROFILE.keys(), n=1, cutoff=0.92)
        if best:
            mapping[col] = COL_ALIASES_PROFILE[best[0]]
    return mapping

def _product_columns(df: pd.DataFrame, colmap: dict) -> list:
    prod_cols = [c for c, canon in colmap.items() if canon == "product_name"]
    if not prod_cols:
        prod_cols = [c for c in df.columns if "product" in str(c).lower() or ("name" in str(c).lower() and "product" in str(c).lower())]
        if not prod_cols:
            prod_cols = [c for c in df.columns if "name" in str(c).lower()]
    return prod_cols

def _best_product_row(df: pd.DataFrame, product_query: str, prod_cols: list) -> Optional[int]:
    if not prod_cols:
        return None
    # exact/ci first
    for c in prod_cols:
        mask = df[c].astype(str).str.strip().str.lower() == product_query.strip().lower()
        if mask.any():
            return mask[mask].index[0]
    # fuzzy across all candidate values
    candidates: List[str] = []
    for c in prod_cols:
        vals = df[c].dropna().astype(str).tolist()
        candidates.extend(vals)
    best = difflib.get_close_matches(product_query, candidates, n=1, cutoff=0.7)
    if best:
        target = best[0]
        for c in prod_cols:
            mask = df[c].astype(str).str.strip() == target
            if mask.any():
                return mask[mask].index[0]
    # contains fallback
    ql = product_query.lower()
    for c in prod_cols:
        mask = df[c].astype(str).str.lower().str.contains(re.escape(ql), na=False)
        if mask.any():
            return mask[mask].index[0]
    return None

def build_product_profile_from_df(df: pd.DataFrame, product_query: str) -> Optional[str]:
    if df is None or df.empty:
        return None
    colmap = _map_columns_profile(df)
    prod_cols = _product_columns(df, colmap)
    if not prod_cols:
        return None

    ridx = _best_product_row(df, product_query, prod_cols)
    if ridx is None:
        return None

    row = df.loc[ridx]
    canon: Dict[str, str] = {}
    for orig, key in colmap.items():
        val = row.get(orig, None)
        if pd.notna(val) and str(val).strip() and key:
            canon[key] = str(val).strip()

    if "product_name" not in canon:
        for c in prod_cols:
            val = row.get(c, None)
            if pd.notna(val) and str(val).strip():
                canon["product_name"] = str(val).strip()
                break

    if not canon:
        return None

    lines: List[str] = []
    title = canon.get("product_name", product_query)
    lines.append(f"*• Product:* {title}")
    ordered_keys = [k for k in PREFERRED_PROFILE_ORDER if k in canon] + [k for k in canon.keys() if k not in PREFERRED_PROFILE_ORDER]

    label_map = {
        "other_names": "Also known as",
        "description": "Description",
        "product_manager": "Product Manager",
        "product_mgmt_vp": "Product Mgmt VP",
        "development_vp": "Development VP",
        "support_owner": "Support Owner",
        "second_line_owner": "2nd/1st Line Owner",
        "support_director": "Support Director",
        "support_planner": "Support Planner",
        "pillar": "Pillar",
        "mmt_name": "MMT",
        "mmt_url": "MMT URL",
    }
    for k in ordered_keys:
        if k == "product_name":
            continue
        label = label_map.get(k, k.replace("_", " ").title())
        lines.append(f"*• {label}:* {canon[k]}")
    return "\n".join(lines)

# -----------------------------------------------------------------------------
# Q&A over Excel (STRICT)
# -----------------------------------------------------------------------------

def answer_from_excel_super_dynamic(df: pd.DataFrame, question: str) -> str:
    import re as _re
    import difflib as _difflib
    import pandas as _pd

    def __pp_normalize_dashes(s: str) -> str:
        return s.replace("–", "-").replace("—", "-")

    # Local copies for isolation
    COL_ALIASES_PROFILE = {
        "product name": "product_name",
        "product": "product_name",
        "name": "product_name",

        "other names this product is called (e.g. old names people still use)": "other_names",
        "aka": "other_names",
        "aliases": "other_names",

        "brief description of what product does for a client": "description",

        "product manager": "product_manager",
        "product mgmt vp": "product_mgmt_vp",
        "vp product mgmt": "product_mgmt_vp",
        "development vp": "development_vp",
        "support owner": "support_owner",
        "2nd line or first line owner": "second_line_owner",
        "support director": "support_director",
        "support planner": "support_planner",
        "pillar": "pillar",
        "mmt name (meeting where the cross functional team's meet and vote)": "mmt_name",
        "mmt name https://w3.ibm.com/w3publisher/software-support-community/new-page": "mmt_url",
    }
    PREFERRED_PROFILE_ORDER = [
        "product_name",
        "other_names",
        "description",
        "product_manager",
        "product_mgmt_vp",
        "development_vp",
        "support_owner",
        "second_line_owner",
        "support_director",
        "support_planner",
        "pillar",
        "mmt_name",
        "mmt_url",
    ]

    def __pp_map_columns_profile(df_: _pd.DataFrame) -> dict:
        mapping = {}
        for col in df_.columns:
            key = str(col).strip().lower()
            try:
                key = _clean_text(key)
            except Exception:
                key = _re.sub(r"[-_/]+", " ", key)
                key = _re.sub(r"\s+", " ", key).strip()
            if key in COL_ALIASES_PROFILE:
                mapping[col] = COL_ALIASES_PROFILE[key]
                continue
            best = _difflib.get_close_matches(key, COL_ALIASES_PROFILE.keys(), n=1, cutoff=0.92)
            if best:
                mapping[col] = COL_ALIASES_PROFILE[best[0]]
        return mapping

    def __pp_product_columns(df_: _pd.DataFrame, colmap: dict) -> list:
        prod_cols = [c for c, canon in colmap.items() if canon == "product_name"]
        if not prod_cols:
            prod_cols = [c for c in df_.columns if "product" in str(c).lower() or ("name" in str(c).lower() and "product" in str(c).lower())]
            if not prod_cols:
                prod_cols = [c for c in df_.columns if "name" in str(c).lower()]
        return prod_cols

    def __pp_best_product_row(df_: _pd.DataFrame, product_query: str, prod_cols: list):
        if not prod_cols:
            return None
        for c in prod_cols:
            mask = df_[c].astype(str).str.strip().str.lower() == product_query.strip().lower()
            if mask.any():
                return mask[mask].index[0]
        candidates = []
        for c in prod_cols:
            vals = df_[c].dropna().astype(str).tolist()
            candidates.extend(vals)
        best = _difflib.get_close_matches(product_query, candidates, n=1, cutoff=0.7)
        if best:
            target = best[0]
            for c in prod_cols:
                mask = df_[c].astype(str).str.strip() == target
                if mask.any():
                    return mask[mask].index[0]
        ql = product_query.lower()
        for c in prod_cols:
            mask = df_[c].astype(str).str.lower().str.contains(_re.escape(ql), na=False)
            if mask.any():
                return mask[mask].index[0]
        return None

    def __pp_build_product_profile_from_df(df_: _pd.DataFrame, product_query: str):
        if df_ is None or df_.empty:
            return None
        colmap = __pp_map_columns_profile(df_)
        prod_cols = __pp_product_columns(df_, colmap)
        if not prod_cols:
            return None

        ridx = __pp_best_product_row(df_, product_query, prod_cols)
        if ridx is None:
            return None

        row = df_.loc[ridx]
        canon = {}
        for orig, key in colmap.items():
            val = row.get(orig, None)
            if _pd.notna(val) and str(val).strip() and key:
                canon[key] = str(val).strip()

        if "product_name" not in canon:
            for c in prod_cols:
                val = row.get(c, None)
                if _pd.notna(val) and str(val).strip():
                    canon["product_name"] = str(val).strip()
                    break

        if not canon:
            return None

        lines = []
        title = canon.get("product_name", product_query)
        lines.append(f"*• Product:* {title}")
        ordered_keys = [k for k in PREFERRED_PROFILE_ORDER if k in canon] + [k for k in canon.keys() if k not in PREFERRED_PROFILE_ORDER]
        label_map = {
            "other_names": "Also known as",
            "description": "Description",
            "product_manager": "Product Manager",
            "product_mgmt_vp": "Product Mgmt VP",
            "development_vp": "Development VP",
            "support_owner": "Support Owner",
            "second_line_owner": "2nd/1st Line Owner",
            "support_director": "Support Director",
            "support_planner": "Support Planner",
            "pillar": "Pillar",
            "mmt_name": "MMT",
            "mmt_url": "MMT URL",
        }
        for k in ordered_keys:
            if k == "product_name":
                continue
            label = label_map.get(k, k.replace("_", " ").title())
            lines.append(f"*• {label}:* {canon[k]}")
        return "\n".join(lines)

    # -------------------------------------------------
    # main logic
    # -------------------------------------------------
    q = __pp_normalize_dashes(question).lower().strip()

    ROLE_HINTS = ["marketing head", "head of marketing", "marketing lead"]
    known_cols_lc = [c.lower() for c in df.columns]
    if any(h in q for h in ROLE_HINTS):
        if not any("marketing" in c for c in known_cols_lc):
            return NOT_FOUND_MSG

    if q.startswith("full_product_profile::"):
        product_query = q.split("::", 1)[1].strip().strip('"\'')

        product_query = _re.sub(r"\s+", " ", product_query)
        prof = __pp_build_product_profile_from_df(df, product_query)
        if prof:
            return prof

    m_direct = _re.match(r"^-\s*(?:g\s+)?product\s+(.+)$", q, _re.IGNORECASE)
    if m_direct:
        product_query = m_direct.group(1).strip().strip('"\'')

        product_query = _re.sub(r"\s+", " ", product_query)
        prof = __pp_build_product_profile_from_df(df, product_query)
        if prof:
            return prof

    m = _re.search(r"last\s+(\d+)\s+(rows|data|entries|records|activities|tasks|items)", q)
    if m:
        n = int(m.group(1))
        last_rows = df.tail(n)
        if len(last_rows) == 1:
            row = last_rows.iloc[0]
            return "\n".join(f"{col}: {row[col]}" for col in df.columns)
        try:
            return last_rows.to_markdown(index=False)
        except Exception:
            return last_rows.to_string(index=False)

    m = _re.search(
        r"(?:^|\s)(?:what|which|who)\s+(?:is\s+)?(?:the\s+)?(.+?)\s+(?:of|for|does)\s+(?:the\s+)?(?:product\s+)?[\"']?(.+?)[\"']?(?:\s+belong to)?[\?\.]?$",
        q
    )
    if m:
        col_query = _clean_text(m.group(1))
        entity    = _clean_entity(m.group(2))

        role_col, role_score = resolve_role_column(df.columns, col_query)
        if not role_col:
            return NOT_FOUND_MSG
        if re.search(r"\b(product|name)\b", role_col, re.I):
            return NOT_FOUND_MSG

        product_cols = [c for c in df.columns if "product" in c.lower() or ("name" in c.lower() and "product" in c.lower())]
        candidates = _pd.DataFrame()
        search_cols = product_cols or list(df.columns)

        for pc in search_cols:
            mask = _two_way_contains(df[pc], entity)
            if mask.any():
                part = df[mask].copy()
                part["_match_col"] = pc
                candidates = _pd.concat([candidates, part], axis=0)

        if not candidates.empty:
            row = candidates.iloc[0]
            match_col = row["_match_col"] if "_match_col" in row else search_cols[0]
            matched_label = row[match_col] if match_col in row.index else "the item"

            if (role_col in df.columns) and _pd.notna(row.get(role_col, None)) and str(row.get(role_col)).strip():
                return f"The {role_col} of {matched_label} is: {row[role_col]}"
            return NOT_FOUND_MSG

        entity_cols = [c for c in df.columns if any(word in c.lower() for word in
                        ["product","name","owner","person","employee","user","activity","item","project"])]
        for entity_col in entity_cols:
            mask = _two_way_contains(df[entity_col], entity)
            matches = df[mask]
            if not matches.empty:
                row = matches.iloc[0]
                if (role_col in df.columns) and _pd.notna(row.get(role_col, None)) and str(row.get(role_col)).strip():
                    return f"The {role_col} of {entity} is: {row[role_col]}"
                return NOT_FOUND_MSG

    m = _re.search(r"(?:who\s+is|details\s+of)\s+(.+?)[\?\.]?$", q)
    if m:
        entity = m.group(1).strip()
        all_matches = _pd.DataFrame()
        for col in df.columns:
            mask = df[col].astype(str).str.lower().str.contains(entity.lower())
            if mask.any():
                all_matches = _pd.concat([all_matches, df[mask]])

        all_matches = all_matches.drop_duplicates()

        if not all_matches.empty:
            role_info = {}
            for col in df.columns:
                if any(word in col.lower() for word in ["owner", "manager", "lead", "vp", "director", "planner"]):
                    count = (all_matches[col].astype(str).str.lower() == entity.lower()).sum()
                    if count > 0:
                        role_info[col] = count

            if not role_info:
                return NOT_FOUND_MSG

            primary_role = max(role_info.items(), key=lambda x: x[1])[0]
            mask = df[primary_role].astype(str).str.lower() == entity.lower()
            person_rows = df[mask]

            if person_rows.empty:
                return NOT_FOUND_MSG

            item_col = None
            for col in df.columns:
                if any(word in col.lower() for word in ["product", "activity", "task", "item", "project"]) and "name" in col.lower():
                    item_col = col
                    break
            if not item_col:
                for col in df.columns:
                    if "name" in col.lower():
                        item_col = col
                        break

            items: List[str] = []
            if item_col and item_col in person_rows.columns:
                items = [s for s in person_rows[item_col].dropna().astype(str).tolist() if s and s.lower() != 'nan']

            entity_title = entity.title()
            response = f"{entity_title} is the {primary_role} of {len(person_rows)} entr{'y' if len(person_rows)==1 else 'ies'}."
            if items:
                if len(items) == 1:
                    response += f" They are: {items[0]}."
                else:
                    response += f" They are: {', '.join(items[:-1])}, and {items[-1]}."
            return response

        return NOT_FOUND_MSG

    m = _re.search(r"show\s+(\d+)(?:st|nd|rd|th)?\s+row", q)
    if m:
        row_num = int(m.group(1)) - 1
        if 0 <= row_num < len(df):
            row = df.iloc[row_num]
            return "\n".join(f"{col}: {row[col]}" for col in df.columns if _pd.notna(row[col]) and str(row[col]) != 'nan')
        else:
            return f"Row {row_num + 1} does not exist. The data has {len(df)} rows."

    if any(phrase in q for phrase in ["show all", "list all", "display all", "all rows", "all entries", "show everything", "full data", "complete data"]):
        return "I'm not able to retrieve all data. Please refine your query for specific information.[example: 'show me the last 3 rows', 'show me the first row'. etc.]"

    m = _re.search(r"(?:products?\s+under|list\s+(?:the\s+)?products?\s+under)\s+(.+)", q)
    if m:
        alias = m.group(1).strip().lower()
        colmap = __pp_map_columns_profile(df)
        alias_cols = [c for c, canon in colmap.items() if canon == "other_names"]
        prod_cols = __pp_product_columns(df, colmap)

        if alias_cols and prod_cols:
            matches = df[df[alias_cols[0]].astype(str).str.lower().str.contains(alias, na=False)]
            if not matches.empty:
                products = matches[prod_cols[0]].dropna().astype(str).tolist()
                if products:
                    return f"Products under {alias.upper()}: " + ", ".join(products)

        return f"I couldn't find any products under {alias.upper()} in this sheet."

    return NOT_FOUND_MSG
